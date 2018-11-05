#-*- encoding: utf8 -*-
import docx
file=docx.Document("D:\\pyS\\iTN8600-SQ4_A_SYSTEM_7.6.32_20180224.docx")
a=u'啊啊啊'
print(a)
print(u'段落数:'+str(len(file.paragraphs)))

for para in file.paragraphs:
    print(para.text.decode('utf-8'))

for i in range(len(file.paragraphs)):
    print(u'第'+str(i)+u'段的内容是：'+(file.paragraphs[i].text).decode('utf-8'))